from docx import Document
import base64
from io import BytesIO
import plotly.express as px


def doctable(df, title, hcols, fcols):
    for i, col in enumerate(hcols):
        df[col].fillna(method='ffill',inplace=True)
        
    # df fillna
    df = df.fillna('')
    tb=df.groupby(hcols,sort=False).size().reset_index(name='Sum')
    # create picture
    fig = px.pie(tb, values='Sum', names=hcols[0], title='Graph')
    fig.update_traces(textposition='inside', textinfo='label+value+percent')

    df3=df.groupby(hcols)[fcols].agg(list).reset_index()
    document = Document()
    document.add_heading(title, 0)

    document.add_heading('Table', 1)
    
    table = document.add_table(tb.shape[0]+1, tb.shape[1])

    
    for j in range(tb.shape[-1]):
        table.cell(0,j).text = tb.columns[j]

    for i in range(tb.shape[0]):
        for j in range(tb.shape[-1]):
            table.cell(i+1,j).text = str(tb.values[i,j])
    
    document.add_paragraph()
    document.add_paragraph('Total: '+str(len(df)))
   
    document.add_page_break()
    document.add_heading('Listing', 1)
    
    count = 1
    for index, row in df3.iterrows():
        for i, col in enumerate(hcols):
            head = col + ' ' + str(row[col])
            document.add_heading(head, 2)
        lastlv = 2
        item_num = len(row[fcols[0]])
        for j in range(item_num):
            document.add_heading('No: ' + str(count), lastlv+1)
            count += 1
            for k, colm in enumerate(fcols):
                document.add_heading(colm, lastlv+2)
                document.add_paragraph(str(row[colm][j]))
        document.add_page_break()
    return document,fig

def download_word_link(object_to_download, download_filename, download_link_text):
    # create buffer
    buffer = BytesIO()
    object_to_download.save(buffer)
    buffer.seek(0)
    # create base64
    b64 = base64.b64encode(buffer.read()).decode()
    return f'<a href="data:file/txt;base64,{b64}" download="{download_filename}">{download_link_text}</a>'


def download_svg_link(fig, download_filename, download_link_text):
    # create buffer
    buffer = BytesIO()
    # convert to svg
    fig.write_image(buffer, format='svg')
    buffer.seek(0)
    # create base64
    # retrieve svg
    svg = buffer.getvalue()
    b64 = base64.b64encode(svg).decode()
    # create html
    return f'<a href="data:image/svg+xml;base64,{b64}" download="{download_filename}">{download_link_text}</a>'